#!/usr/bin/env python3
"""Rebuild the Kingridge (Cawthra) Inc. rescission form as a clean, professional,
Google-Docs-friendly .docx. Import-safe features only: tables with cell shading &
borders, merged header rows, colored/spaced text, paragraph borders for rule lines.
No text boxes, no floating shapes, no nested tables (all degrade on Google Docs)."""

from docx import Document
from docx.shared import Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- palette ----------
NAVY   = "1E3A5F"; STEEL = "2E5A88"; BAND = "EEF2F7"; BOXBG = "F5F7FA"
BORDER = "D5DCE4"; RULE  = "AEB9C7"; MUTED = "5B6673"; INK  = "232323"
WHITE  = "FFFFFF"; SUBTLE = "C9D6E8"
HEAD_FONT = "Georgia"; BODY_FONT = "Arial"
CW = 9936  # content width in twips (letter 12240 - margins 1152*2)

# ---------- helpers ----------
def _sub(parent, tag, **attrs):
    el = OxmlElement(tag)
    for k, v in attrs.items():
        el.set(qn(k), str(v))
    parent.append(el)
    return el

def style_run(run, size=None, color=INK, bold=False, italic=False,
              font=BODY_FONT, spacing=None, caps=False):
    if caps and run.text:
        run.text = run.text.upper()
    run.font.name = font
    if size is not None:
        run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = _sub(rpr, 'w:rFonts')
    for a in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rfonts.set(qn(a), font)
    if spacing is not None:
        sp = rpr.find(qn('w:spacing'))
        if sp is None:
            sp = OxmlElement('w:spacing'); rpr.append(sp)
        sp.set(qn('w:val'), str(spacing))
    return run

def add_run(p, text, **kw):
    return style_run(p.add_run(text), **kw)

def para_format(p, before=0, after=0, line=None, align=None, keep_next=False):
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line; pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if align is not None:
        p.alignment = align
    if keep_next:
        pf.keep_with_next = True
    return p

def set_borders(pr, edges, tag='w:tcBorders'):
    ex = pr.find(qn(tag))
    if ex is not None:
        pr.remove(ex)
    b = OxmlElement(tag)
    for name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        if name in edges:
            spec = edges[name]; e = _sub(b, f'w:{name}')
            if spec is None:
                e.set(qn('w:val'), 'nil')
            else:
                size, color = spec
                e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), str(size))
                e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
    pr.append(b)

def shade(pr, color):
    ex = pr.find(qn('w:shd'))
    if ex is not None:
        pr.remove(ex)
    _sub(pr, 'w:shd', **{'w:val': 'clear', 'w:color': 'auto', 'w:fill': color})

def cell_margins(cell, top=60, bottom=60, left=130, right=130):
    tcPr = cell._tc.get_or_add_tcPr()
    m = tcPr.find(qn('w:tcMar'))
    if m is None:
        m = _sub(tcPr, 'w:tcMar')
    for name, val in (('top', top), ('bottom', bottom),
                      ('start', left), ('end', right),
                      ('left', left), ('right', right)):
        e = m.find(qn(f'w:{name}'))
        if e is None:
            e = _sub(m, f'w:{name}')
        e.set(qn('w:w'), str(val)); e.set(qn('w:type'), 'dxa')

def cell_shade(cell, color):
    shade(cell._tc.get_or_add_tcPr(), color)

def cell_borders(cell, edges):
    set_borders(cell._tc.get_or_add_tcPr(), edges)

def cell_valign(cell, val):
    cell.vertical_alignment = val

def set_col_widths(table, widths):
    table.autofit = False; table.allow_autofit = False
    tblPr = table._tbl.tblPr
    layout = tblPr.find(qn('w:tblLayout'))
    if layout is None:
        layout = _sub(tblPr, 'w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    grid = table._tbl.find(qn('w:tblGrid'))
    if grid is not None:
        table._tbl.remove(grid)
    grid = OxmlElement('w:tblGrid')
    table._tbl.insert(list(table._tbl).index(tblPr) + 1, grid)
    for w in widths:
        _sub(grid, 'w:gridCol', **{'w:w': str(w)})
    for row in table.rows:
        for i, c in enumerate(row.cells):
            if i < len(widths):
                c.width = Twips(widths[i])
                tcPr = c._tc.get_or_add_tcPr()
                tcw = tcPr.find(qn('w:tcW'))
                if tcw is None:
                    tcw = _sub(tcPr, 'w:tcW')
                tcw.set(qn('w:w'), str(widths[i])); tcw.set(qn('w:type'), 'dxa')

def table_total_width(table, total=CW):
    tblPr = table._tbl.tblPr
    w = tblPr.find(qn('w:tblW'))
    if w is None:
        w = _sub(tblPr, 'w:tblW')
    w.set(qn('w:w'), str(total)); w.set(qn('w:type'), 'dxa')
    ind = tblPr.find(qn('w:tblInd'))
    if ind is None:
        ind = _sub(tblPr, 'w:tblInd')
    ind.set(qn('w:w'), '0'); ind.set(qn('w:type'), 'dxa')

def clear_table_borders(table):
    set_borders(table._tbl.tblPr,
                {e: None for e in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']},
                tag='w:tblBorders')

def table_borders(table, edges):
    set_borders(table._tbl.tblPr, edges, tag='w:tblBorders')

def row_height(row, twips, exact=False):
    trPr = row._tr.get_or_add_trPr()
    h = trPr.find(qn('w:trHeight'))
    if h is None:
        h = _sub(trPr, 'w:trHeight')
    h.set(qn('w:val'), str(twips))
    h.set(qn('w:hRule'), 'exact' if exact else 'atLeast')

def new_table(cols_widths, n_rows):
    t = doc.add_table(rows=n_rows, cols=len(cols_widths))
    table_total_width(t); set_col_widths(t, cols_widths); clear_table_borders(t)
    return t

# ---------- document ----------
doc = Document()
sec = doc.sections[0]
sec.page_height = Twips(15840); sec.page_width = Twips(12240)
sec.top_margin = Twips(1008); sec.bottom_margin = Twips(1008)
sec.left_margin = Twips(1152); sec.right_margin = Twips(1152)
sec.header_distance = Twips(600); sec.footer_distance = Twips(600)

normal = doc.styles['Normal']
normal.font.name = BODY_FONT; normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
rpr = normal.element.get_or_add_rPr()
rf = rpr.find(qn('w:rFonts'))
if rf is None:
    rf = _sub(rpr, 'w:rFonts')
for a in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
    rf.set(qn(a), BODY_FONT)
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.line_spacing = 1.0

def spacer(pts):
    p = doc.add_paragraph()
    para_format(p, after=0, line=1.0)
    p.add_run("").font.size = Pt(pts)
    return p

def section_header_cell(hc, title):
    cell_shade(hc, BAND)
    cell_borders(hc, {'left': (28, NAVY), 'bottom': (8, BORDER)})
    cell_margins(hc, top=82, bottom=82, left=170, right=150)
    hc.text = ""
    p = hc.paragraphs[0]; para_format(p, after=0, line=1.0, keep_next=True)
    add_run(p, title, size=10.5, color=NAVY, bold=True, spacing=40, caps=True)

def merge_row(table, r, ncols):
    hc = table.cell(r, 0)
    for j in range(1, ncols):
        hc = hc.merge(table.cell(r, j))
    return hc

# ============================================================= HEADER (one table)
head = new_table([6600, 3336], 3)
# row0 masthead
m = merge_row(head, 0, 2)
cell_shade(m, NAVY); cell_margins(m, top=195, bottom=205, left=250, right=250)
m.text = ""
p = m.paragraphs[0]; para_format(p, after=2, line=1.0)
add_run(p, "KINGRIDGE (CAWTHRA) INC.", size=9, color=SUBTLE, bold=True, spacing=60, caps=True)
p = m.add_paragraph(); para_format(p, after=0, line=1.0)
add_run(p, "Rescission Form", size=28, color=WHITE, bold=True, font=HEAD_FONT)
p = m.add_paragraph(); para_format(p, before=6, after=0, line=1.0)
add_run(p, "Cancellation of Agreement of Purchase & Sale", size=11, color=SUBTLE, italic=True)
# row1 accent strip
a = merge_row(head, 1, 2)
cell_shade(a, STEEL); cell_margins(a, 0, 0, 0, 0); a.text = ""
a.paragraphs[0].add_run("").font.size = Pt(2)
row_height(head.rows[1], 58, exact=True)
# row2 info (property | rescission date)
il, ir = head.rows[2].cells
for cell in (il, ir):
    cell_shade(cell, BAND); cell_margins(cell, top=115, bottom=120, left=210, right=210)
    cell_valign(cell, WD_CELL_VERTICAL_ALIGNMENT.CENTER); cell.text = ""
cell_borders(il, {'right': (18, WHITE)})
p = il.paragraphs[0]; para_format(p, after=1, line=1.0)
add_run(p, "PROPERTY", size=8, color=MUTED, bold=True, spacing=50, caps=True)
p = il.add_paragraph(); para_format(p, after=0, line=1.0)
add_run(p, "3111 Cawthra Road", size=13, color=NAVY, bold=True, font=HEAD_FONT)
p = ir.paragraphs[0]; para_format(p, after=2, line=1.0)
add_run(p, "RESCISSION DATE", size=8, color=MUTED, bold=True, spacing=50, caps=True)
p = ir.add_paragraph(); para_format(p, before=4, after=0, line=1.0)
add_run(p, "  ", size=11)
set_borders(p._element.get_or_add_pPr(), {'bottom': (6, RULE)}, tag='w:pBdr')

# ---------- field-section table (header row + label/value rows) ----------
def field_section(title, rows, label_w=3050):
    spacer(6)
    n = 1 + len(rows)
    t = new_table([label_w, CW - label_w], n)
    section_header_cell(merge_row(t, 0, 2), title)
    table_borders(t, {'insideH': (6, BORDER), 'bottom': (6, BORDER)})
    for i, (label, value) in enumerate(rows, start=1):
        row_height(t.rows[i], 348)
        lc, vc = t.rows[i].cells
        for cell in (lc, vc):
            cell_valign(cell, WD_CELL_VERTICAL_ALIGNMENT.CENTER)
            cell_margins(cell, top=88, bottom=88, left=150, right=150)
            cell.text = ""
        lp = lc.paragraphs[0]; para_format(lp, after=0, line=1.0)
        add_run(lp, label, size=10, color=NAVY, bold=True)
        vp = vc.paragraphs[0]; para_format(vp, after=0, line=1.0)
        add_run(vp, value if value else "", size=10.5, color=INK)
    return t

field_section("Notice of Rescission", [
    ("To", "Kingridge (Cawthra) Inc."),
    ("From (Purchaser)", ""),
    ("Purchaser Name (1)", ""),
    ("Purchaser Name (2)", ""),
])
field_section("Purchase & Property Details", [
    ("Date of Sale", ""),
    ("Date of Acknowledgement", ""),
    ("Block / Unit #", ""),
    ("Model", ""),
    ("Sq. Ft.", ""),
])

# ---------- Declaration ----------
spacer(6)
dt = new_table([CW], 2)
section_header_cell(dt.cell(0, 0), "Declaration")
dc = dt.cell(1, 0)
cell_shade(dc, BAND); cell_borders(dc, {'left': (22, STEEL)})
cell_margins(dc, top=140, bottom=150, left=200, right=220); dc.text = ""
p = dc.paragraphs[0]; para_format(p, after=0, line=1.18)
add_run(p, "Re: Cancellation of Purchase.  ", size=11, color=NAVY, bold=True)
add_run(p, "The undersigned wish to rescind the Agreement of Purchase & Sale "
           "for the above-mentioned property.", size=11, color=INK, italic=True)

# ---------- Signatures ----------
spacer(6)
w3 = CW // 3
st = new_table([w3, w3, CW - 2 * w3], 3)
section_header_cell(merge_row(st, 0, 3), "Signatures")
sig_labels = ["Purchaser Signature", "Witness Signature", "Date"]
for r in (1, 2):
    row_height(st.rows[r], 640)
    for i, cell in enumerate(st.rows[r].cells):
        cell_valign(cell, WD_CELL_VERTICAL_ALIGNMENT.BOTTOM)
        cell_margins(cell, top=40, bottom=45, left=95, right=180); cell.text = ""
        line = cell.paragraphs[0]; para_format(line, after=3, line=1.0)
        add_run(line, "", size=10)
        set_borders(line._element.get_or_add_pPr(), {'bottom': (6, RULE)}, tag='w:pBdr')
        lbl = cell.add_paragraph(); para_format(lbl, after=0, line=1.0)
        add_run(lbl, sig_labels[i], size=8.5, color=MUTED, bold=True, spacing=20, caps=True)

# ---------- Reason for Head Office Records ----------
spacer(6)
rt = new_table([CW], 4)
section_header_cell(rt.cell(0, 0), "Reason for Head Office Records")
def reason_cell(idx, build):
    c = rt.cell(idx, 0)
    cell_margins(c, top=70, bottom=70, left=150, right=150); c.text = ""
    build(c.paragraphs[0])
def chk(p, text, tail=""):
    para_format(p, after=0, line=1.15)
    add_run(p, "☐  ", size=12, color=NAVY)
    add_run(p, text, size=10.5, color=INK)
    if tail:
        add_run(p, tail, size=10.5, color=MUTED)
reason_cell(1, lambda p: chk(p, "Could not qualify for a mortgage"))
reason_cell(2, lambda p: chk(p, "Other:  ",
            "______________________________________________________"))
def deposit(p):
    para_format(p, after=0, line=1.2)
    add_run(p, "Original Initial Deposit Cheque of ", size=10.5, color=INK)
    add_run(p, "$ __________ ", size=10.5, color=NAVY, bold=True)
    add_run(p, "and ", size=10.5, color=INK)
    add_run(p, "__________ ", size=10.5, color=NAVY, bold=True)
    add_run(p, "post-dated cheque(s) returned.", size=10.5, color=INK)
reason_cell(3, deposit)

# ---------- For Administrative Use Only (single flat table, box look) ----------
spacer(9)
label_w = 5600
at = new_table([label_w, CW - label_w], 4)
table_borders(at, {'top': (8, BORDER), 'bottom': (8, BORDER),
                   'left': (8, BORDER), 'right': (8, BORDER)})
hc = merge_row(at, 0, 2)
cell_shade(hc, BAND); cell_borders(hc, {'left': (28, NAVY), 'bottom': (8, BORDER)})
cell_margins(hc, top=95, bottom=95, left=170, right=150); hc.text = ""
p = hc.paragraphs[0]; para_format(p, after=0, line=1.0)
add_run(p, "For Administrative Use Only", size=10, color=NAVY, bold=True, spacing=45, caps=True)
admin_items = ["Received Condo Docs & APS", "Cheques Returned", "Cancelled in System"]
for i, item in enumerate(admin_items, start=1):
    row_height(at.rows[i], 360)
    lc, vc = at.rows[i].cells
    for cell in (lc, vc):
        cell_shade(cell, BOXBG); cell_valign(cell, WD_CELL_VERTICAL_ALIGNMENT.BOTTOM)
        cell.text = ""
    cell_margins(lc, top=80, bottom=75, left=170, right=60)
    cell_margins(vc, top=80, bottom=75, left=30, right=170)
    lp = lc.paragraphs[0]; para_format(lp, after=0, line=1.0)
    add_run(lp, "☐  ", size=12, color=NAVY)
    add_run(lp, item, size=10, color=INK)
    vp = vc.paragraphs[0]; para_format(vp, after=0, line=1.0)
    add_run(vp, "Date / Initials", size=7.5, color=MUTED, caps=True, spacing=20)
    cell_borders(vc, {'bottom': (6, RULE)})

# ---------- Footer ----------
footer = sec.footer; footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
para_format(fp, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
set_borders(fp._element.get_or_add_pPr(), {'top': (6, BORDER)}, tag='w:pBdr')
add_run(fp, "Kingridge (Cawthra) Inc.", size=8, color=MUTED, bold=True, spacing=20)
add_run(fp, "    •    Rescission Form  —  3111 Cawthra Road", size=8, color=MUTED)

import os
OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                   "Rescission-Form-3111-Cawthra-Road.docx")
doc.save(OUT)
print("Saved:", OUT)
print("Top-level tables:", len(doc.tables))
